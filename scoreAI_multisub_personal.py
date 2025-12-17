# scoreAI_multisub_personal.py

# This script is adapted from Wardell et al., 2021 (see below)
# edited to work with Hennessy et al. scoring procedures and team management (40+ RAs, weekly assignments)

# takes in scored word documents, creates 
# 1) a spreadsheet for each memory, with 1 row per scorer (for IRR chats)
# 2) an itial scored sheet for all memories


###############################################################
##                                                           ##
##  scoreAI                                                  ##
##  -------                                                  ##
##                                                           ##
##  Python script for processing transcribed and scored      ##
##  autobiographical memory narratives.                      ##
##                                                           ##
##  Current public version: build 10 [20200511]              ##
##                                                           ##
##                                                           ##
##  If you use this software, please cite:                   ##
##  * Wardell, V., Esposito, C. L., Madan, C. R., & Palombo, ##
##      D. J. (2021). Semi-automated transcription and       ##
##      scoring of autobiographical memory narratives.       ##
##      Behavior Research Methods, 53, 507-517.              ##
##                                                           ##
##  For documentation and further information, see:          ##
##     https://github.com/cMadan/scoreAI                     ##
##                                                           ##
###############################################################

## EDITED BY SARAH HENNESSY MARCH 7, 2023 for the SEMPRE study

# main changes:

## 1. nMemories per file always = 1
## 2. Output csv filename is always the coded sub ID
## 3. New spreadsheet for each participant, 2 scorers per file 
## 4. Iterates through files but groups based on ID so that we can achieve ^ 
## 5. change in tags, slightly
## 6. Change in header, slightly
## 7. Output is an xlsx to account for multiple tabs


from docx import Document #python-docx package
import os
from os import listdir
from os.path import isfile, join
import pandas as pd
from datetime import date
from warnings import warn
import re
import subprocess
import shutil

red= "\033[31m"

###############################################################
## SECTION 1: Config
# folder with memory docs

set = input("What set are you on? (p,1-5): \n")

pathIn = 'scoreAI/Data/input/set_%s' %(set)
pathOut = 'scoreAI/Data/output/set_%s' %(set)


#First, pull out all of the scored docs from the personal folders and put them in the Input folder
personalIn = 'scoreAI/Data/Personal Folders'

if not os.path.exists(pathIn):
    os.mkdir(pathIn)
    print("Directory", pathIn, "created")
else:
    print("Directory", pathIn, "already exists")


if not os.path.exists(pathOut):
    os.mkdir(pathOut)
    print("Directory", pathOut, "created")
else:
    print("Directory", pathOut, "already exists")


personalIn = 'scoreAI/Data/Personal Folders'


for RAfolder in os.listdir(personalIn):
    print(RAfolder)
    RApath = personalIn + "/" + RAfolder
    try:
        setpath = RApath + "/Scored/Set %s" %(set)
        print(setpath)
        for doc in os.listdir(setpath):
            if doc.endswith('.docx'):
                print(doc)
                docpath = setpath +  "/" + doc
                shutil.copy(docpath, pathIn)
                print("copied")

    except:
        #print("not in file, moving on?")
        continue




print("copied all the folders")
# how many memories are we expecting per document?
nMemories = 1


###############################################################
## SECTION 3: Expert settings
# memory scoring codes
tags = ['Int_EV','Int_PL','Int_TM','Int_PER','Int_EMO','Ext_EV','Ext_PL','Ext_TM','Ext_PER','Ext_EMO','Ext_SEM','Ext_REP','Ext_OTH']

# file lists
#docs = [f for f in listdir(pathIn) if isfile(join(pathIn, f)) if not f.startswith('.')]
docs = [f for f in listdir(pathIn) if isfile(join(pathIn, f)) if not f.startswith('.') and not f.startswith('~')]


# min length (characters) of a valid memory response,
# used to parse doc and skip empty runs
# only used by getResponse, but not getPara
minresponselength = 25


###############################################################
## SECTION 4: Define re-usable functions
# seek next paragraph that includes specific text string
def seekPara(para,string):
    search = True
    while search:
        para += 1
        try:
            # strip removes trailing spaces
            #text = d.paragraphs[para].runs[0].text.strip()
            # above approach is too brittle to idiosyncrasies of Word XML formatting
            text = getPara(para).strip()
            # check if line matches specific search string
            compare = text == string
            if compare:
                search = False # found it!
            elif para == len(runcount)-1:
                # didn't find the string,
                # shouldn't ever happen
                para = False
                search = False # abort
        except:
            # do nothing, will fail often
            False
    return para

# pull paragraph text
def getPara(para):
    response = ''
    for run in d.paragraphs[para].runs:
        response += run.text
    response = response.strip()
    return response

# pull paragraph text for a transcribed response
def getResponse(para):
    search = True
    while search:
       minlength = runcount[para][1]>minresponselength
       if minlength:
           search = False # found it!
       else:
           para += 1
    # para found, now compile text across runs
    response = getPara(para)
    return response

# get counts for each tag
def countTag(response):
    tagCount = []
    for tag in tags:
        count = response.count(tag)
        tagCount.append(count)
    return tagCount


def getTitle():
    search = True
    titlelist = []
    para = 0
    for p in d.paragraphs:
        para += 1
        text = ''
        for runs in p.runs:
            text += runs.text
        # adjust for potential leading space
        text = text.strip()

        # in the docs, sometimes the next character was either a hyphen or an endash
        if text[0:6] == 'Title:':
            titleitem = text[7:]
            if len(titleitem) == 0:
                titleitem = "NONE"
            titlelist.append(titleitem)
            #titlelist.append(text[7:])


    return titlelist

def getSpecificity():
    search = True
    speclist = []
    para = 0
    ticker = 0
    is_spec = 0
    for p in d.paragraphs:
        # print("I HAVE THIS MANY PARAS")
        # print(len(d.paragraphs))
        ticker = ticker + 1

        para += 1

        text = ''
        for runs in p.runs:
            text += runs.text
        # adjust for potential leading space
        text = text.strip()

        # in the docs, sometimes the next character was either a hyphen or an endash
    #    print("ticker is: %d" %(ticker))
        if ticker < len(d.paragraphs):
        #    print("Going into normal loop")
            if 'Specificity' in text:
                is_spec = 1
                if text[0:12] == 'Specificity:':
                    specitem = text[13:]
                    if len(specitem) == 0:
                        specitem = "NONE"
                    speclist.append(specitem)
        else:
       
            if is_spec == 1:
                
                print("I found specificity  before...")
                # print(speclist)
            else:
                # print("I didnt find specificity before..")
                if text[0:12] == 'Specificity:':
                    specitem = text[13:]
                    if len(specitem) == 0:
                        specitem = "NONE"
                    speclist.append(specitem)
                    # print(speclist)
                else:
                    # print("I didnt find specificity this time either..")
                    # print("NO SPEC")
                    specitem = "NO SPECIFICITY"
                    speclist.append(specitem)
                    # print(speclist)





    return speclist


def getTimePeriod():
    # get episodic richness values for each memory
    search = True
    timelist = []
    # identify when ER is stated
    para = 0
    for p in d.paragraphs:
        para += 1
        text = ''
        for runs in p.runs:
            text += runs.text
        # adjust for potential leading space
        text = text.strip()
        if text[0:6] == 'Time P':
            timeitem = text[13:]
            if len(timeitem) == 0:
                timeitem = "NONE"
            timelist.append(timeitem)
    return timelist

def getPlace():
    # get episodic richness values for each memory
    search = True
    placelist = []
    # identify when ER is stated
    para = 0
    for p in d.paragraphs:
        para += 1
        text = ''
        for runs in p.runs:
            text += runs.text
        # adjust for potential leading space
        text = text.strip()

        # in the docs, sometimes the next character was either a hyphen or an endash
        if text[0:9] == 'Place Loc':
            placeitem = text[20:]
            if len(placeitem) == 0:
                placeitem = "NONE"

            # get the specific ER values
            placelist.append(placeitem)
    return placelist

def getTimeLoc():
    # get episodic richness values for each memory
    search = True
    timeloclist = []
    # identify when ER is stated
    para = 0
    for p in d.paragraphs:
        para += 1
        text = ''
        for runs in p.runs:
            text += runs.text
        # adjust for potential leading space
        text = text.strip()

        # in the docs, sometimes the next character was either a hyphen or an endash
        if text[0:18] == 'Time Localization:':
            timelocitem = text[19:]
            if len(timelocitem) == 0:
                timelocitem = "NONE"
            # get the specific ER values
            timeloclist.append(timelocitem)

    return timeloclist

def getPer():
    # get episodic richness values for each memory
    search = True
    perlist = []
    # identify when ER is stated
    para = 0
    for p in d.paragraphs:
        para += 1
        text = ''
        for runs in p.runs:
            text += runs.text
        # adjust for potential leading space
        text = text.strip()

        # in the docs, sometimes the next character was either a hyphen or an endash
        if text[0:20] == 'Perceptual Richness:':
            peritem = text[21:]
            if len(peritem) == 0:
                peritem = "NONE"
            perlist.append(peritem)

    return perlist

def getEmo():
    # get episodic richness values for each memory
    search = True
    emolist = []
    # identify when ER is stated
    para = 0
    for p in d.paragraphs:
        para += 1
        text = ''
        for runs in p.runs:
            text += runs.text
        # adjust for potential leading space
        text = text.strip()

        # in the docs, sometimes the next character was either a hyphen or an endash
        if text[0:9] == 'Emotions:':
            emoitem = text[10:]
            if len(emoitem) == 0:
                emoitem = "NONE"

            emolist.append(emoitem)

    return emolist

def getTimeInt():
    # get episodic richness values for each memory
    search = True
    timeintlist = []
    # identify when ER is stated
    para = 0
    for p in d.paragraphs:
        para += 1
        text = ''
        for runs in p.runs:
            text += runs.text
        # adjust for potential leading space
        text = text.strip()

        # in the docs, sometimes the next character was either a hyphen or an endash
        if text[0:17] == 'Time Integration:':
            intitem = text[18:]
            if len(intitem) == 0:
                intitem = "NONE"
            timeintlist.append(intitem)

    return timeintlist

def getEpi():
    # get episodic richness values for each memory
    search = True
    epilist = []
    # identify when ER is stated
    para = 0
    for p in d.paragraphs:
        para += 1
        text = ''
        for runs in p.runs:
            text += runs.text
        # adjust for potential leading space
        text = text.strip()

        # in the docs, sometimes the next character was either a hyphen or an endash
        if text[0:18] == 'Episodic Richness:':

            richitem = text[19:]
            if len(richitem) == 0:
                richitem = "NONE"
            epilist.append(richitem)

    return epilist

def getmem():
    # get episodic richness values for each memory
    search = True
    memlist = []
    # identify when ER is stated
    para = 0
    for p in d.paragraphs:
        para += 1
        text = ''
        for runs in p.runs:
            text += runs.text
        # adjust for potential leading space
        text = text.strip()

        # in the docs, sometimes the next character was either a hyphen or an endash
        if text[0:18] == 'Memory about song:':
            memitem = text[19:]
            if len(memitem) == 0:
                memitem = "NONE"
            memlist.append(memitem)

    return memlist



###############################################################
## SECTION 5: Process the data
# cycle through each doc in input folder

subIDlist = []

for doc in docs:

    print('Processing '+doc)
    try:
    # load document

        d = Document(join(pathIn,doc))

        ## parse document
        # enumerate paragraph runs
        runcount = [len(d.paragraphs[p].runs) for p in range(len(d.paragraphs))]
        textcount = []
        for p in d.paragraphs:
            text = ''
            for runs in p.runs:
                text += runs.text
            textcount.append(len(text))
        runcount = list(zip(runcount,textcount))

        # get the Participant ID
        cover = d.tables[0].rows[0].cells[0].text
        cover = cover.split('\n')
    #    subID = cover[1].replace('Participant ID: ','').strip()

# just make it the name of the file
        subID = doc.split("_")[0]

       # if subID in subIDlist:



        scorer = cover[2].replace('Scorer: ','').strip()

        # get title (main event)
        titlelist = getTitle()
        speclist = getSpecificity()
        timelist = getTimePeriod()
        placelist = getPlace()
        timeloclist = getTimeLoc()
        perlist = getPer()
        emolist =getEmo()
        timeintlist = getTimeInt()
        epilist = getEpi()
        memlist = getmem()
        if len(speclist) == 0:
            print(red+ "I FOUND IT")

        # get start and end paras for each memory
        para = 0
        paraM = []
        for M in list(range(1,nMemories+1)):
            # find para number for Memory M
            para = seekPara(para,'Memory '+str(M))
            paraM.append(para)
            print(str(M)+'...', end = '')
        if len(paraM) != nMemories:
            warn('WARNING: In %s, Number of memories found does not match number of memories expected. %g memories found, expected %g.' % (doc,len(paraM),nMemories))

        # add last run number as end value for para intervals
        paraM.append(len(runcount))

        # get text and counts for each memory
        responseCounts = []
        for M in list(range(1,nMemories+1)):
            response = ''
            for para in list(range(paraM[M-1],paraM[M])):
                response += getPara(para)
                response += '|'
            counts = countTag(response)
            responseCounts.append(counts)

        ## output data
        # init dict first
        data_sub = {
            'ParticipantID': [subID] * nMemories, # repeat subID several times, like repmat
           'Scorer': [scorer],
            'Memory': list(range(1,nMemories+1))
        }


        # restructure tags to have list of each tag and count for each memory
        counts = dict(zip(tags, list(map(list,zip(*responseCounts))) ))


        # concatenate dicts, to merge in the tag counts
        data_sub.update(counts)

        # concat in the episodic richness values

        data_sub.update({'Title':titlelist})
        data_sub.update({'Specificity':speclist})
        data_sub.update({'TimePeriod':timelist})
        data_sub.update({'PlaceLocalization':placelist})
        data_sub.update({'TimeLocalization':timeloclist})
        data_sub.update({'PerceptualRichness':perlist})
        data_sub.update({'Emotions':emolist})
        data_sub.update({'TimeIntegration':timeintlist})
        data_sub.update({'EpisodicRichness':epilist})
        data_sub.update({'MemoryAboutSong':memlist})


        # merge data to common record, or create it if not exists
        try:
            data_all = { key:data_all.get(key,[])+data_sub.get(key,[]) for key in data_all.keys() }
        except NameError:
            data_all = data_sub
        print('done.')

    except Exception as e:
        print(e)


print("data all: ")
print(data_all)
###############################################################
## SECTION 6: Output the counts
# convert data dict into a dataframe

# if you get an error here, you probably have a file that doesn't have a specificity part
# you can figure out which it is by  adding this line.
print("padding...")

# get the maximum length of the lists
max_len = max(len(lst) for lst in data_all.values())

# pad the lists with None if they're shorter than the maximum length
for key in data_all.keys():
    data_all[key] = data_all[key] + [1000] * (max_len - len(data_all[key]))


data_all = pd.DataFrame(data_all)

#print(data_all)
print('got data DataFrame')


###############################################################

#make one big spreadsheet

fOut = 'set%s.xlsx' %(set)

data_all.to_excel(join(pathOut,fOut),index=False)


checked = input("GREAT. now go look at scoreAI/Data/output/set_x/setx.xlsx and make sure that it has no Nones or 1000s. Type y when you have completed that.: \n" )

###############################################################
## fin

if checked == "y":

    #ORGANIZE THE FILES....

    #write to file, one file per sub ID.
    def extractIDs(set,pathOut):

        # Read original Excel file
      
        df = pd.read_excel(pathOut + '/set%s.xlsx' %(set))

        # Get unique subIDs
        subIDs = df['ParticipantID'].unique()


        # Iterate over subIDs and create a new Excel file for each
        for subID in subIDs:
            # Filter out rows for current subID
            sub_df = df[df['ParticipantID'] == subID]
            print("sub ID is %s" %(subID))


            mask = sub_df.notnull().any(axis=1)   # boolean mask for rows with at least one non-null value

            first_row_index = sub_df.first_valid_index()
            first_row = sub_df.loc[first_row_index]
            RA_name1 = first_row['Scorer']
            print(RA_name1)

            second_row_index = sub_df.loc[first_row_index+1:].first_valid_index()



            if second_row_index is not None:
                second_row = sub_df.loc[second_row_index]
                RA_name2 = second_row['Scorer']
                print(RA_name2)
            else:
                print("There is no second row with data.")
                RA_name2 = 'None'

           
            setname = "Set "+ set
            dest_folder1 = os.path.join(pathOut, RA_name1, setname)


            # make the destination folder if it doesn't already exist
            if not os.path.exists(dest_folder1):
                os.makedirs(dest_folder1)


            sub_df.to_excel(dest_folder1 +'/{}.xlsx'.format(subID), index=False)

            dest_folder2 = os.path.join(pathOut, RA_name2, setname)


            # make the destination folder if it doesn't already exist
            if not os.path.exists(dest_folder2):
                os.makedirs(dest_folder2)


            sub_df.to_excel(dest_folder2 +'/{}.xlsx'.format(subID), index=False)
    

    extractIDs(set,pathOut)
    print("Done Scoring!")


    for RA in os.listdir(pathOut):
        if RA.endswith('.xlsx') == 0:
            print(RA)
            RApathset = pathOut + "/" + RA + "/Set %s" %(set)
            personalRApath = personalIn + "/" + RA + "/" + "Spreadsheets/Set %s" %(set)
            if not os.path.exists(personalRApath):
                os.mkdir(personalRApath)
                print("Directory", personalRApath, "created")
            for sheet in os.listdir(RApathset):
                if sheet.endswith('.xlsx'):
                    sheetpath = RApathset+ "/" + sheet
                    shutil.copy(sheetpath, personalRApath)
                    print("copied")

    print("Done Copying!")
